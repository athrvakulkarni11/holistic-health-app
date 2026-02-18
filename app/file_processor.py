"""
File Processing Service — Extracts text and biomarker data from uploaded files.
Supports PDF, images (via OCR-like LLM vision), CSV, TXT, and DOCX formats.
Uses Groq LLM to intelligently extract biomarker values from unstructured text.
"""
import os
import json
import re
import csv
import io
from groq import Groq
from app.config import GROQ_API_KEY, GROQ_MODEL, UPLOAD_DIR


class FileProcessorService:
    def __init__(self):
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)

        self.groq_client = None
        if GROQ_API_KEY:
            self.groq_client = Groq(api_key=GROQ_API_KEY)

    def save_file(self, filename: str, content: bytes) -> str:
        """Save uploaded file and return the path."""
        filepath = os.path.join(UPLOAD_DIR, filename)
        with open(filepath, "wb") as f:
            f.write(content)
        return filepath

    def extract_text(self, filepath: str) -> str:
        """Extract text from various file formats."""
        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".pdf":
            return self._extract_pdf(filepath)
        elif ext in (".png", ".jpg", ".jpeg"):
            return self._extract_image(filepath)
        elif ext == ".csv":
            return self._extract_csv(filepath)
        elif ext == ".txt":
            return self._extract_txt(filepath)
        elif ext == ".docx":
            return self._extract_docx(filepath)
        else:
            return f"Unsupported file format: {ext}"

    def _extract_pdf(self, filepath: str) -> str:
        """Extract text from PDF using PyPDF2 or pdfplumber."""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # Also try extracting tables
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text_parts.append(" | ".join([str(cell) if cell else "" for cell in row]))

            return "\n".join(text_parts) if text_parts else "Could not extract text from PDF."
        except ImportError:
            # Fallback to PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n".join(text_parts) if text_parts else "Could not extract text from PDF."
            except ImportError:
                return "PDF parsing libraries not available. Please install pdfplumber or PyPDF2."
            except Exception as e:
                return f"Error reading PDF: {str(e)}"
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    def _extract_image(self, filepath: str) -> str:
        """For images, we describe the limitation and attempt OCR if available."""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(filepath)
            text = pytesseract.image_to_string(img)
            return text if text.strip() else "Could not extract text from image via OCR."
        except ImportError:
            return (
                "Image uploaded. OCR libraries (pytesseract, Pillow) not available. "
                "Please type or paste the values from your lab report."
            )
        except Exception as e:
            return f"Error processing image: {str(e)}"

    def _extract_csv(self, filepath: str) -> str:
        """Extract data from CSV file."""
        try:
            with open(filepath, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = list(reader)
            lines = []
            for row in rows:
                lines.append(" | ".join(row))
            return "\n".join(lines)
        except Exception as e:
            return f"Error reading CSV: {str(e)}"

    def _extract_txt(self, filepath: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {str(e)}"

    def _extract_docx(self, filepath: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(cells))
            return "\n".join(text_parts) if text_parts else "Could not extract text from DOCX."
        except ImportError:
            return "python-docx library not available. Please install it: pip install python-docx"
        except Exception as e:
            return f"Error reading DOCX: {str(e)}"

    def extract_biomarkers_from_text(self, raw_text: str) -> dict:
        """
        Use Groq LLM to intelligently extract biomarker values from unstructured text.
        Returns a dict with profile and biomarker data ready for analysis.
        """
        if not self.groq_client:
            return {"error": "LLM not available for biomarker extraction."}

        system_prompt = """You are a medical data extraction AI. Your task is to extract biomarker values and patient profile information from lab report text.

Extract the following fields if available. Return ONLY valid JSON:

{
  "profile": {
    "age": null,
    "gender": null,
    "height": null,
    "weight": null
  },
  "biomarkers": {
    "hemoglobin": null,
    "rbc_count": null,
    "ferritin": null,
    "vitamin_b12": null,
    "vitamin_d": null,
    "fasting_glucose": null,
    "hba1c": null,
    "total_cholesterol": null,
    "ldl": null,
    "hdl": null,
    "triglycerides": null,
    "hs_crp": null,
    "tsh": null,
    "sgpt_alt": null
  },
  "additional_biomarkers": {},
  "extraction_notes": "Any observations about the data quality or missing values"
}

Rules:
- Only populate fields where you find clear numerical values
- Convert units if necessary to match the expected units (g/dL for hemoglobin, etc.)
- If a value seems to be in a different unit, note the conversion in extraction_notes
- Set null for any values not found in the text
- Include any other biomarkers found in additional_biomarkers as key:value pairs
- Be VERY accurate — do not guess or hallucinate values"""

        user_prompt = f"""Extract biomarker values from this lab report text:

{raw_text}

Return only valid JSON with the extracted values."""

        try:
            response = self.groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.1,
                max_tokens=2048,
                response_format={"type": "json_object"},
            )
            content = response.choices[0].message.content
            return json.loads(content)
        except json.JSONDecodeError:
            return {"error": "Could not parse extracted data", "raw": content}
        except Exception as e:
            return {"error": f"Extraction failed: {str(e)}"}

    def process_uploaded_file(self, filename: str, content: bytes) -> dict:
        """
        Full pipeline: save file → extract text → extract biomarkers.
        Returns structured data ready for analysis.
        """
        filepath = self.save_file(filename, content)
        raw_text = self.extract_text(filepath)
        
        # Privacy Guard Rail: Delete file after extraction
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except Exception as e:
            print(f"Error deleting file {filepath}: {e}")

        if raw_text.startswith("Error") or raw_text.startswith("Unsupported"):
            return {"success": False, "error": raw_text, "raw_text": ""}

        extracted = self.extract_biomarkers_from_text(raw_text)

        return {
            "success": True,
            "raw_text": raw_text,
            "extracted_data": extracted,
            "filename": filename,
        }
